import os
import re
from pathlib import Path
from typing import List, Dict, Any
from pypdf import PdfReader
from backend.config import CHUNK_SIZE, CHUNK_OVERLAP

try:
    import docx
except ImportError:
    docx = None

class DocumentProcessor:
    """Handles text extraction and semantic chunking with metadata from cybersecurity documents (PDF, DOCX, DOC, MD, TXT)."""

    def __init__(self, chunk_size: int = CHUNK_SIZE, chunk_overlap: int = CHUNK_OVERLAP):
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap

    def process_file(self, file_path: Path) -> List[Dict[str, Any]]:
        """Processes a single file (PDF, DOCX, DOC, MD, or TXT) and returns chunk dictionaries with metadata."""
        suffix = file_path.suffix.lower()
        chunks = []

        if suffix == ".pdf":
            chunks = self._process_pdf(file_path)
        elif suffix in [".docx", ".doc"]:
            chunks = self._process_docx(file_path)
        elif suffix in [".md", ".txt"]:
            chunks = self._process_text_file(file_path)
        else:
            print(f"[DocumentProcessor] Unsupported file type: {file_path.name}")
        
        return chunks

    def _process_pdf(self, file_path: Path) -> List[Dict[str, Any]]:
        """Extracts text from PDF page-by-page and splits into chunks."""
        chunks = []
        filename = file_path.name
        doc_title = file_path.stem.replace("_", " ").title()

        try:
            reader = PdfReader(file_path)
            for page_num, page in enumerate(reader.pages, start=1):
                text = page.extract_text() or ""
                text = self._clean_text(text)
                if not text:
                    continue

                page_chunks = self._split_text_into_chunks(text)
                for idx, chunk_text in enumerate(page_chunks):
                    chunk_id = f"{file_path.stem}_p{page_num}_c{idx}"
                    chunks.append({
                        "id": chunk_id,
                        "text": chunk_text,
                        "metadata": {
                            "source": filename,
                            "title": doc_title,
                            "page": page_num,
                            "chunk_index": idx,
                            "file_type": "pdf"
                        }
                    })
        except Exception as e:
            print(f"[DocumentProcessor] Error reading PDF {filename}: {e}")

        return chunks

    def _process_docx(self, file_path: Path) -> List[Dict[str, Any]]:
        """Extracts text from Word documents (.docx / .doc) and splits into chunks."""
        chunks = []
        filename = file_path.name
        doc_title = file_path.stem.replace("_", " ").title()

        try:
            if docx is not None:
                doc = docx.Document(file_path)
                full_text = []
                for para in doc.paragraphs:
                    if para.text.strip():
                        full_text.append(para.text.strip())
                
                for table in doc.tables:
                    for row in table.rows:
                        row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                        if row_text:
                            full_text.append(" | ".join(row_text))

                text_content = "\n\n".join(full_text)
            else:
                # Fallback text reading
                with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                    text_content = f.read()

            cleaned_content = self._clean_text(text_content)
            raw_chunks = self._split_text_into_chunks(cleaned_content)

            for idx, chunk_text in enumerate(raw_chunks):
                chunk_id = f"{file_path.stem}_doc_c{idx}"
                chunks.append({
                    "id": chunk_id,
                    "text": chunk_text,
                    "metadata": {
                        "source": filename,
                        "title": doc_title,
                        "page": f"Section {idx+1}",
                        "chunk_index": idx,
                        "file_type": file_path.suffix[1:]
                    }
                })
        except Exception as e:
            print(f"[DocumentProcessor] Error reading Word document {filename}: {e}")
            # Plaintext fallback attempt
            chunks = self._process_text_file(file_path)

        return chunks

    def _process_text_file(self, file_path: Path) -> List[Dict[str, Any]]:
        """Extracts text from Markdown or TXT files and splits into chunks."""
        filename = file_path.name
        doc_title = file_path.stem.replace("_", " ").title()

        try:
            with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                content = f.read()

            cleaned_content = self._clean_text(content)
            raw_chunks = self._split_text_into_chunks(cleaned_content)

            chunks = []
            for idx, chunk_text in enumerate(raw_chunks):
                heading_match = re.search(r'(?:^|\n)#+\s+(.+)', chunk_text)
                section = heading_match.group(1).strip() if heading_match else f"Section {idx+1}"
                
                chunk_id = f"{file_path.stem}_c{idx}"
                chunks.append({
                    "id": chunk_id,
                    "text": chunk_text,
                    "metadata": {
                        "source": filename,
                        "title": doc_title,
                        "page": section,
                        "chunk_index": idx,
                        "file_type": file_path.suffix[1:]
                    }
                })
            return chunks
        except Exception as e:
            print(f"[DocumentProcessor] Error reading text file {filename}: {e}")
            return []

    def _clean_text(self, text: str) -> str:
        """Removes extraneous whitespace while preserving structure."""
        text = re.sub(r'\r\n', '\n', text)
        text = re.sub(r'[ \t]+', ' ', text)
        text = re.sub(r'\n{3,}', '\n\n', text)
        return text.strip()

    def _split_text_into_chunks(self, text: str) -> List[str]:
        """Splits long text into overlapping chunks based on character/token boundaries."""
        if len(text) <= self.chunk_size:
            return [text]

        paragraphs = text.split("\n\n")
        chunks = []
        current_chunk = []
        current_length = 0

        for para in paragraphs:
            para = para.strip()
            if not para:
                continue

            para_len = len(para)
            if current_length + para_len > self.chunk_size and current_chunk:
                combined = "\n\n".join(current_chunk)
                chunks.append(combined)
                
                overlap_text = combined[-self.chunk_overlap:] if len(combined) > self.chunk_overlap else combined
                current_chunk = [overlap_text, para]
                current_length = len(overlap_text) + para_len
            else:
                current_chunk.append(para)
                current_length += para_len

        if current_chunk:
            chunks.append("\n\n".join(current_chunk))

        return chunks
